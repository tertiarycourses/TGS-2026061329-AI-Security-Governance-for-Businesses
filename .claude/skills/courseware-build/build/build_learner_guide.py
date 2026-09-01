#!/usr/bin/env python3
"""Generate the AI Security Governance for Businesses Learner Guide as BOTH a Markdown mirror
(LG-*.md at repo root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt body, one section
per lab (Objective · Goal · What you'll build · Step-by-step · Test it · Worked guidance), plus
setup, framework reference and glossary. All content is driven by course_data + the domain data
files, keeping the LG 100% aligned with the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

def _asset(name):
    """Resolve a brand asset: the course's own courseware/assets first, then the skill's."""
    for base in (os.path.join(REPO,"courseware","assets"), ASSETS):
        q=os.path.join(base,name)
        if os.path.exists(q): return q
    return None


LAB_DIRS={
 1:"lab-01-ai-inventory",2:"lab-02-threat-landscape",3:"lab-03-lethal-trifecta",4:"lab-04-business-case",
 5:"lab-05-nist-rmf-gap",6:"lab-06-policy-set",7:"lab-07-operating-model",8:"lab-08-pdpa-application",
 9:"lab-09-data-governance",10:"lab-10-red-team",11:"lab-11-deployment-monitoring",12:"lab-12-incident-response",
 13:"lab-13-agentic-controls",14:"lab-14-csa-agentic-profile",15:"lab-15-risk-assessment",16:"lab-16-capstone-roadmap"}

# Per-lab expert guidance shown in the LG (not on the slides).
GUIDANCE={
 1:["Work the three files in order — IT assets, then procurement, then the staff survey. Each layer finds systems the previous one missed, and that progression is itself the lesson.",
    "The AI indicator column in the IT extract is an automatic flag and it is not reliable. Read the description column yourself: 'Document OCR Pipeline' is an AI system even where the flag disagrees.",
    "Expect roughly a dozen AI systems once you merge all three sources, of which at least three appear only in the staff survey. Those three are your shadow-AI findings.",
    "A system with no named owner is a finding in its own right. Do not quietly assign it to IT to make the register look tidy — record the gap."],
 2:["Run each simulator module before you fill in its row. The point is to observe the mechanism, not to recall the definition.",
    "In the SQL injection module, read the live query display carefully. Seeing WHERE user='admin' AND pass='' OR '1'='1' makes clear why the fix is parameterisation — the input became part of the instruction.",
    "That is exactly the reasoning to carry into prompt injection, with one crucial difference: a prompt has no parameterised form. Data and instructions arrive in the same channel and there is no equivalent of a bound variable. Write that difference in your worksheet.",
    "The data-leakage estimator is worth pausing on: notice which single control moves the risk score most, and ask whether your organisation actually has it."],
 3:["Do not skip to the redesign. Complete all three legs of the trifecta first — the redesign is only defensible once you can show the full exposure.",
    "Leg 2 is the one people under-count. Anything an outsider can influence is untrusted: customer e-mail, uploaded PDFs, web search results, and even a third-party API response.",
    "For the attack path, follow the EchoLeak shape: attacker plants text where retrieval will find it, the agent reads it as instruction, the agent uses a permitted tool, data leaves. No exploit is required at any step.",
    "The strongest redesigns split the agent: a retrieval agent that reads untrusted content but holds no account data and no outbound channel, feeding a summary to a second agent that never sees raw external text. Adding an output filter to the original design is the weaker answer, and you should be able to say why."],
 4:["Use one scenario, costed properly, rather than a list. The prompt-injection leak from a customer-facing agent is the most defensible because you analysed the design in Lab 3.",
    "Annual expected loss = annual likelihood x total impact. Take likelihood from sector-incident-data.csv and build impact from the four cost categories. Show the arithmetic so a CFO can challenge an assumption rather than the whole number.",
    "State your assumptions explicitly on the page. A number without visible assumptions gets dismissed; a number with them gets debated, which is what you want.",
    "Prepare for 'why not just buy a tool?'. The answer: a tool enforces policy — with no inventory, no owner and no policy, there is nothing for it to enforce, and you will have bought a dashboard."],
 5:["Score only on the evidence in the pack. If the interview does not support a 2, it is a 1, however capable the organisation feels.",
    "GOVERN 1.6 (AI inventory) is a 0 here — the CMDB excludes SaaS AI features and the CRO guesses at 'three or four' systems. That single gap invalidates most downstream controls, which is why it belongs near the top of your register.",
    "Note the trap in Extract D: model retraining and prompt changes are treated as 'configuration' and bypass CAB entirely. That is a change-control gap that no amount of infrastructure review will catch.",
    "MEASURE will score lowest. No AI system is tested for bias, robustness or jailbreak resistance — the Head of Digital Channels confirms 200 accuracy questions and nothing else.",
    "Prioritise by risk reduction per unit of effort. Turning on logging is a small effort closing a High risk, and it should outrank a large effort closing a Medium one."],
 6:["Write for the person who must follow it, not for the auditor who will file it. The test at the end of the lab is the real quality bar: could a colleague read it and know what to do on Monday?",
    "Be explicit about scope. The most common failure is a policy that never says whether it covers M365 Copilot, a personal ChatGPT account, or an AI feature switched on inside a SaaS tool you already own.",
    "Prohibited-use rules must be concrete enough to detect. 'Use AI responsibly' is unenforceable; 'do not enter customer personal data into any tool not on the approved list' can be evidenced and trained.",
    "Commit to an approval turnaround in days. Without a service level, staff route around the process, and your policy has quietly created more shadow AI than it prevented.",
    "The MGF mapping is not decoration — it is how you show a Singapore regulator or a client that your policy set aligns with the national framework."],
 7:["Reuse before you create. The CRO has explicitly refused a new committee unless you show existing forums cannot absorb the remit, and that constraint reflects real organisational life.",
    "Decision rights are the section that matters. Write which decisions this committee takes, which escalate to the ERC and which reach the board — a committee that can only advise cannot govern.",
    "Apply the single-A rule strictly. If two roles are Accountable for one control, each will assume the other has it, and you will discover the gap during an incident.",
    "The stress test is the real assessment: walk the NovaAssist incident through your model. Detection, the decision to disable, and the PDPA notification must each land on a named role with no gaps and no overlaps."],
 8:["Read the scenario for what the system does to a person, not for what technology it uses. That is what drives the PDPA analysis.",
    "Scenario 1 is the clearest Business Improvement Exception case — improving an existing product with an existing customer's data. Record the basis in writing; the exception does not remove the Accountability Obligation.",
    "Scenario 2 materially affects an individual, so the notification and human-oversight requirements are strongest here. Note the postal-code feature: it can act as a proxy for ethnicity or income and create discrimination risk even where the field is lawful to use.",
    "Scenario 3 is a controls question rather than a lawful-basis question. The risk is staff pasting customer data into prompts, so the answer is classification rules, input filtering, logging and retention limits.",
    "Scenario 4: the vendor is a data intermediary carrying the Protection and Retention Obligations. On 'anonymised learnings', ask what exactly is anonymised — a model trained on your customers' data is not obviously anonymous, and this belongs in the contract, not in an e-mail."],
 9:["Classify before you minimise, and minimise before you de-identify. Doing it in that order removes most of the work, because a dropped field needs no protection at all.",
    "Ask the minimisation question honestly for each field: does churn prediction need it? Name, NRIC, e-mail, mobile and address plainly do not. Religion and ethnicity not only fail the test but create discrimination risk if the model learns from them.",
    "Postal code is the interesting case. It has predictive value and it is a strong quasi-identifier and a proxy for income. Generalising to a district is usually the right call — record your reasoning either way.",
    "In the Cryptography Toolkit, encrypt an identifier with AES and then decrypt it back. That round trip is the point: encryption is reversible, so an encrypted NRIC remains personal data to whoever holds the key.",
    "Use ECDSA signing for dataset integrity — sign the approved training set so you can later prove it was not altered between approval and training. This is the control auditors ask for and few organisations have.",
    "Never describe the result as 'anonymised' if you hashed identifiers. It is pseudonymised, and PDPA obligations continue to apply."],
 10:["Read the FauxBank disclaimer and take it seriously. These techniques are for authorised testing only — unauthorised access to a computer is a criminal offence under the Singapore Computer Misuse Act.",
     "Record evidence as you go, not afterwards. For each finding write the input, the observed response and what it proves about the control that failed. Findings without evidence do not survive review.",
     "IDOR deserves particular attention in an AI context: an agent iterating over identifiers will walk an IDOR at a speed no human tester would, turning a moderate finding into a mass-disclosure event.",
     "Compare your manual findings against the simulated scanner deliberately. Each will find things the other missed, and that comparison is your argument for why automated assurance alone is not assurance.",
     "The five AI-specific test cases are the part no scanner produces. Test indirect injection especially — instructions hidden in retrieved content, including text made invisible with white-on-white formatting.",
     "A recommendation without conditions is not assurance. If you recommend GO, state precisely what must be true first."],
 11:["Design the gate around evidence, not intentions. Each item should be a document or a test result someone can produce, not an assurance that it was considered.",
     "Decide the hard case in advance: what happens when the business wants to launch with an item unsatisfied. Deciding this during a launch means the gate loses.",
     "Every exception needs an expiry date and a compensating control. An exception without an expiry is a permanent silent gap that nobody revisits.",
     "In the log extract, work the four anomalies: a burst of ~48 rapid calls from one user, a run of 11 consecutive refusals, an injection flag followed within seconds by an outbound send, and a cluster of out-of-hours activity around 03:14.",
     "The injection-then-send sequence is the critical one — it is exactly the Lab 3 attack path appearing in telemetry. Your alerting must catch the sequence, not just the individual flag.",
     "A metric with no threshold is a dashboard. Give every metric a number and a named response action, and measure your time to disable rather than estimating it."],
 12:["Stay in your role. The exercise tests whether the operating model you designed produces a named decision-maker at every step, and stepping outside your role hides the gaps.",
     "Containment first, analysis second. Decide whether to disable fully or restrict outbound tools — restricting send while keeping read may preserve service and stop the harm, and that judgement is the governance skill.",
     "Preserve evidence before anything is reset. Prompts, tool-call logs, model version, agent configuration, approval records and the injected e-mails themselves. Once a system is redeployed, this is gone.",
     "Inject 2 changes the scope from one customer to a systemic failure across 340 runs. If your response does not change with it, revisit it.",
     "On the journalist: agree who speaks, and say only what is established. 'We are investigating a potential issue and will update' is defensible; speculation about numbers is not.",
     "For the notification determination, record the reasoning and not just the conclusion. 214 individuals with names, partial account numbers and balances is the fact pattern to reason from.",
     "The root cause is architectural: private data access plus untrusted content plus an outbound channel, with no injection alerting and no approval gate on send."],
 13:["Classify every tool by consequence before you decide anything. The tool's name tells you nothing; its reversibility and reach tell you everything.",
     "T-08 send_email, T-09 raise_payment, T-10 freeze_account and T-11 partner_api_card are the ones that matter — irreversible or externally visible. Each must be ask or deny, and an allow on any of them needs a very strong justification.",
     "T-12 render_link looks harmless and is not. A rendered markdown link can trigger an automatic fetch to an attacker-controlled URL, which is how EchoLeak exfiltrated data without the user clicking anything.",
     "Write deny rules a policy engine could actually evaluate. 'Deny any outbound send whose body matches an account-number pattern' is enforceable; 'deny unsafe sends' is not.",
     "Be honest in the residual-risk section. A well-built stack still does not stop a legitimate-looking action within the agent's permitted scope, nor a compromised upstream dependency. Naming those two paths is the mark of a real assessment.",
     "Remember the ordering principle: permission enforced before execution is a control; the same rule written in the system prompt is a request that an injection can override."],
 14:["Tier by capability. The IT Helpdesk Agent resets passwords — that is an identity-affecting action, so it is not Tier 1 whatever its description says.",
     "The Fraud Triage Agent freezes customer accounts automatically, roughly 40 times a day, with only after-the-fact weekly review. Work through what a wrong freeze does to a customer, and let that drive the tier.",
     "The Reconciliation Agent is the interesting one: an orchestrator instructs a sub-agent and nobody reviews that instruction. That is the delegation-accountability gap AG-GV.2 exists to close — document the oversight boundary and the lineage back to a named human.",
     "Leave no blank kill-switch field in the registry. If you cannot say how an agent is stopped and how long it takes, you have found a finding, not a formatting problem.",
     "Telemetry needs a baseline before it can have a threshold. Use the Lab 11 log to derive a normal action velocity, then set the alert relative to it.",
     "The Marketing Content Agent decommissioning is deliberately mundane and easy to under-do: revoke credentials, dispose of memory, preserve audit logs for the retention period, and update everything downstream."],
 15:["Scope first and tightly. An assessment of 'our AI' cannot be completed or defended; an assessment of NovaAssist v1.0, its data, its twelve tools and its users can.",
     "Write every risk as cause → event → consequence. If your risk fits in three words it is a cause, not a risk, and it cannot be rated.",
     "Rate inherent risk before controls, and justify every 4 or 5 — those are the ratings that get challenged in committee.",
     "Only credit controls that exist or are formally committed. Crediting a planned control makes residual risk look acceptable while the exposure is unchanged.",
     "Be honest about residual impact. Guardrails and detection reduce likelihood; they rarely reduce impact, because if the data leaves, it has left. Only architectural change reduces impact.",
     "Every residual High risk must be treated or formally accepted by a named executive. 'We will monitor it' is not a treatment.",
     "Check each owner against your Lab 7 operating model. An owner who is not a role in your model is a treatment that will not happen."],
 16:["Consolidate before you plan. Your roadmap must trace to the Lab 5 gap register and the Lab 15 risk assessment — an initiative that closes neither should be cut.",
     "Do not target level 5 across all five domains. Choose deliberately, and be able to say why a domain can sit at 3 for now. A plan that targets everything signals that nothing was prioritised.",
     "Phase 1 is deliberately cheap: inventory with owners, the acceptable-use policy, the committee, and logging. These are low effort and high risk reduction, and they are the dependencies for everything else.",
     "Respect the dependencies literally. You cannot risk-assess systems you have not inventoried, and you cannot enforce a policy you have not issued.",
     "Give every metric a baseline as well as a target. 'Improve inventory coverage' is not measurable; '38% today, 95% in twelve months' is.",
     "On the final slide make one ask. Budget, headcount or mandate — a paper with three asks typically receives none of them.",
     "Rehearse the three challenge questions. The strongest answer to 'is this not just the CISO's job?' is that AI risk spans data, legal, business decisions and security, so it needs a named accountable owner and a cross-functional forum, which is exactly what you designed in Lab 7."],
}


# Screenshots of the browser-based tool each lab uses, embedded in the lab's LG section.
LAB_SHOTS = {
 2:[("tool-threatsim.png","The Cybersecurity Threat Simulator dashboard — the ten threat modules you work through in this lab.")],
 9:[("tool-crypto.png","The Cryptography Toolkit — AES, RSA and ECDSA, used to test protection of the training data.")],
 10:[("tool-fauxbank.png","FauxBank — the pentest training sandbox used for the guided scenarios and the simulated scanner."),
     ("tool-hacklab.png","Hacklab — the simulated terminal used for the reconnaissance and enumeration stages.")],
}

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))
def img(path,caption=""): B.append(("img",path,caption))
def dl(xs): B.append(("dl",xs))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It provides the full step-by-step procedure for all 16 hands-on labs, organised by the four course "
  "topics, together with the reference material and expert guidance you need to complete each lab and "
  "to apply the same method in your own organisation.")
p("Every lab is set in one running case study — NovaBank, a fictional mid-sized Singapore retail bank — "
  "so that your outputs compose into a single coherent governance programme rather than sixteen "
  "disconnected exercises. What you build in Lab 1 is used in Lab 5; what you assess in Lab 5 drives "
  "the policy set in Lab 6, and everything feeds the risk assessment in Lab 15 and the roadmap in Lab 16.")
p("Several labs use browser-based security simulators. These are training environments containing only "
  "fictional data and they generate no real network traffic. The techniques they teach must only ever be "
  "applied to systems you own or have written authorisation to test — unauthorised access to a computer "
  "is a criminal offence under the Singapore Computer Misuse Act and its equivalents elsewhere.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Before You Start — Lab Environment")
h3("What you need")
bullets([
 "A laptop with a modern web browser (Chrome, Edge, Firefox or Safari) and internet access. There is nothing to install.",
 "A spreadsheet application (Excel, Numbers or Google Sheets) for the inventory, RACI, registry and risk-assessment labs.",
 "A word processor or text editor for the policy, report and board-paper labs.",
 "The lab pack from the LMS: one folder per lab containing its mock data files and a printable instruction PDF.",
])
h3("The four browser-based lab tools")
dl([(name,url) for name,url in C.TOOLS])
p("All four are safe to run on a training network. The Hacklab terminal is scripted and produces no real "
  "network traffic; FauxBank is a self-contained sandbox with fictional banking data; the Threat Simulator "
  "and Cryptography Toolkit run entirely in your browser.")
h3("The lab pack layout")
code("labs/\n  lab-01-ai-inventory/\n    lab-01-instructions.pdf     <- printable instructions\n    data/                       <- the mock data for this lab\n      discovery-it-assets.csv\n      discovery-procurement.csv\n      discovery-staff-survey.csv\n  lab-02-threat-landscape/\n  ...\n  lab-16-capstone-roadmap/")
h3("Conventions used in every lab")
bullets([
 "Each lab states an Objective (what it develops), a Goal (the scenario), What you'll build (the artefact you produce) and a Test it (how you know you are done).",
 "Work the steps in order — later steps assume the outputs of earlier ones.",
 "Save every artefact you produce. Later labs consume them, and you may use them in the open-book assessment.",
 "Where a lab asks for a judgement, record the reasoning as well as the conclusion. In governance, the reasoning is the deliverable.",
])
h3("The running case study — NovaBank")
p("NovaBank is a fictional mid-sized Singapore retail bank with about 900 staff. It has a mature "
  "information-security programme, a Data Protection Officer, and a Model Risk Management function that "
  "validates credit and capital models. It has no AI-specific governance at all. It runs a customer-service "
  "agent (NovaAssist), a credit decision engine, a fraud triage agent, an IT helpdesk agent, a reconciliation "
  "agent in pilot, and an unknown quantity of shadow AI. You are its newly appointed AI Governance Lead.")
note("NovaBank, its staff, customers and data are entirely fictional. Any resemblance to a real "
     "organisation or individual is coincidental.")

# ---------------- per-topic, per-lab ----------------
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets([f"{k} — {v}" for k,v in t["concepts"]])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Objective: {a['objective']}")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"])
        h3("Tools and data")
        p(a["services"]+f"   (Lab folder: labs/{LAB_DIRS[a['num']]}/)")
        for _shot,_cap in LAB_SHOTS.get(a["num"],[]):
            img(_shot,_cap)
        h3("Step-by-step")
        steps([(instr,cmd) for instr,cmd in a["steps"]])
        h3("Test it")
        p(a["test"])
        if a["num"] in GUIDANCE:
            h3("Guidance and common pitfalls")
            bullets(GUIDANCE[a["num"]])
        note(f"The mock data and a printable instruction sheet for this lab are in "
             f"labs/{LAB_DIRS[a['num']]}/ (data/ and lab-{a['num']:02d}-instructions.pdf).")
        rule()

# ---------------- reference ----------------
h1("Framework Reference — Ethical Frameworks, Guidelines and Legal Requirements  (K3, A6)")
h2("NIST AI RMF 1.0 — the four functions  (A6)")
dl([
 ("GOVERN","Cultivates a culture of risk management. Policies, processes and procedures (GOVERN 1, including the AI inventory at 1.6 and safe decommissioning at 1.7); accountability structures (GOVERN 2, including executive responsibility at 2.3); diverse teams (GOVERN 3); risk culture (GOVERN 4); stakeholder engagement (GOVERN 5); third-party risk (GOVERN 6). GOVERN is cross-cutting — it is infused through the other three functions, not completed before them."),
 ("MAP","Establishes the context to frame risk: intended purpose and context of use (MAP 1.1), risk tolerances (MAP 1.5), and TEVV considerations (MAP 2.3). Its outcome is enough contextual knowledge to make an initial go/no-go decision."),
 ("MEASURE","Selects metrics and methods, tests and evaluates, and arranges independent internal review (MEASURE 1.3). This is the function most organisations score lowest on."),
 ("MANAGE","Prioritises and treats risk (MANAGE 1.2), documents and accepts residual risk (MANAGE 1.4), monitors after deployment and responds to incidents."),
])
h2("The seven characteristics of trustworthy AI")
bullets([
 "Valid and reliable — it performs as claimed in the conditions it will meet.",
 "Safe — it does not endanger life, health, property or the environment.",
 "Secure and resilient — it withstands adversarial input and recovers from unexpected conditions.",
 "Accountable and transparent — someone is answerable, and information is available to those who need it.",
 "Explainable and interpretable — the mechanism and the meaning of an output can be conveyed appropriately.",
 "Privacy-enhanced — anonymity, confidentiality and individual control are safeguarded.",
 "Fair with harmful bias managed — systemic, computational and human-cognitive bias are identified and managed.",
])
p("These characteristics interact and can trade off against one another. Optimising for one — heavier "
  "privacy protection, say — can degrade another, such as accuracy. Managing that trade-off explicitly, "
  "and recording the decision, is the governance work.")
h2("The 12 NIST generative AI risks (NIST AI 600-1)")
bullets([
 "CBRN information or capabilities — materially lowered barriers to weapons-related information.",
 "Confabulation — confidently stated false content.",
 "Dangerous, violent or hateful content.",
 "Data privacy — leakage or inference of personal data from training data, prompts or outputs.",
 "Environmental impacts — energy and resource cost of training and serving.",
 "Harmful bias and homogenisation — discriminatory outputs and algorithmic monoculture.",
 "Human-AI configuration — over-reliance, poor handoff and rubber-stamped output.",
 "Information integrity — synthetic content degrading trust in what is real.",
 "Information security — prompt injection, model extraction and an expanded attack surface.",
 "Intellectual property — training on, or reproducing, protected material.",
 "Obscene, degrading or abusive content.",
 "Value chain and component integration — third-party models, data and packages you cannot fully inspect.",
])
h2("Singapore — the Model AI Governance Framework for Generative AI")
p("Published by IMDA and the AI Verify Foundation, the framework sets out nine dimensions to be looked "
  "at in totality to foster a trusted ecosystem:")
dl([
 ("Accountability","Allocate responsibility along the AI development chain so end-users have someone answerable."),
 ("Data","Ensure data quality and trusted sources; give business clarity and fair treatment where data is contentious, such as personal data and copyright material."),
 ("Trusted development and deployment","Adopt best practice in development and evaluation, with 'food label'-type transparency on the baseline safety measures taken."),
 ("Incident reporting","Establish structures and processes for incident monitoring, timely notification and remediation."),
 ("Testing and assurance","Use third-party testing for independent verification, and support common standards so results are consistent."),
 ("Security","Recognise that generative AI introduces new threat vectors through the models themselves; adapt existing information-security frameworks."),
 ("Content provenance","Support transparency about where and how content was generated, through techniques such as digital watermarking."),
 ("Safety and alignment R&D","Invest in improving model alignment with human intention and values, cooperating globally."),
 ("AI for public good","Steer development towards broad benefit — access, public-sector adoption, upskilling and sustainability."),
])
p("AI Verify, Singapore's AI governance testing framework and toolkit, is how an organisation "
  "demonstrates through standardised tests that a system behaves as claimed. It produces the evidence "
  "that the trusted-development and testing-and-assurance dimensions call for.")
h2("Singapore — PDPA and the PDPC AI advisory guidelines  (K3)")
p("The PDPC's Advisory Guidelines on the Use of Personal Data in AI Recommendation and Decision Systems "
  "(issued 1 March 2024) explain when personal data may be used to develop and deploy AI systems. They "
  "are not legally binding, but the PDPC will take consistent positions when enforcing the PDPA.")
dl([
 ("Business Improvement Exception","Relevant where the organisation is improving or developing an existing product or service — for example an AI system providing personalised recommendations. It also caters for sharing within a group of related companies, and can apply to bias assessment and to testing the AI system."),
 ("Research Exception","Relevant where the organisation conducts broader commercial research to develop AI systems with public benefit, and covers disclosure for a research purpose subject to conditions."),
 ("Consent and Notification Obligations","Where consent is relied on, notification must be meaningful: individuals should be able to understand what personal data is used and, broadly, how the system uses it to recommend or decide."),
 ("Accountability Obligation","The organisation must be able to SHOW how it discharges its obligations — policies and practices, the written basis for relying on an exception, the measures adopted to protect individuals, and disclosure practices such as model cards and system cards."),
 ("Service Providers","Third-party developers of bespoke AI systems are data intermediaries and carry the Protection and Retention Obligations under the PDPA."),
])
h2("Agentic AI — OWASP ASI Top 10 and the CSA agentic RMF profile")
bullets([
 "ASI01 Goal hijack · ASI02 Tool misuse · ASI03 Identity and authorisation · ASI04 Supply chain · ASI05 Code execution.",
 "ASI06 Memory poisoning · ASI07 RAG poisoning · ASI08 Excessive agency · ASI09 Improper model isolation · ASI10 Operator-facing safety.",
 "AG-GV.1 Autonomy tier classification — four tiers with escalating oversight obligations.",
 "AG-GV.2 Delegation accountability — oversight boundaries, escalation triggers and accountability lineage to a named human.",
 "AG-GV.3 Agent lifecycle registry — a live inventory of agent authorities, tool access and delegation relationships.",
 "AG-MP.1 / AG-MP.2 / AG-MP.3 — tool risk classification, action-consequence analysis and multi-agent topology risk.",
 "AG-MS.1 / AG-MS.2 / AG-MS.3 — behavioural telemetry, autonomy calibration and delegation-chain monitoring.",
 "AG-MG.1 / AG-MG.2 / AG-MG.3 — agentic incident classification, behavioural drift correction and agent decommissioning.",
])
h2("The lethal trifecta")
p("Private data access, exposure to untrusted content, and the ability to communicate externally. An "
  "agent combining all three can be steered by an attacker into disclosing data through entirely "
  "permitted actions — no exploit is required. Removing any one leg is worth more than any filter added "
  "to a design that retains all three. This single pattern explains the majority of publicly reported "
  "agent security incidents in 2025 and 2026.")


h1("Responsible AI Reference — Ethics, Bias, Privacy and Sustainability")
p("The approved skills standard for this course (TSC ICT-INT-0055-1.1, Responsible AI and Generative "
  "AI Practices) assesses responsible-AI knowledge and abilities alongside the security-governance "
  "practice you build in the labs. A system can be perfectly secure and still be unfair, opaque, "
  "privacy-invasive or wasteful — the governance framework has to carry both. This section is the "
  "reference for that material.")
h2("Ethical considerations in AI development (K5)")
dl([
 ("Bias","Address it at every stage: data provenance and representation before training; the fairness measure chosen at design; group-wise testing before release; drift monitoring after it."),
 ("Privacy","Data minimisation and purpose limitation before training; lawful basis under the PDPA; de-identification where feasible; access control; retention limits; and a defined position on deletion requests."),
 ("Transparency","Disclose that AI is in use, what it does and what it cannot do. Meaningful notification where personal data is used, and an explanation where a decision materially affects an individual."),
 ("Accountability","A named human accountable for each AI system — NIST GOVERN 2.1 and 2.3, and the first of the MGF's nine dimensions. In practice: an inventory with an owner per system and exactly one Accountable per control."),
])
p("These considerations trade off against one another. Stronger privacy protection usually costs some "
  "accuracy; more explainability can constrain model choice. Managing and recording that trade-off is "
  "the governance work — pretending it does not exist is the failure. Governance is also proportionate: "
  "risk-tier each system and scale assessment, testing and oversight to the tier.")
h2("How bias enters an AI system, and what it does (K2, A4, A8)")
h3("The five routes bias takes")
bullets([
 "Historical data bias — past decisions are encoded as if they were correct outcomes, so the model reproduces the pattern that produced them.",
 "Representation bias — groups under-represented in the training data are modelled less accurately, so error rates differ systematically between groups.",
 "Proxy discrimination — a feature that is not a protected attribute (postal code, language of a record, booking channel) stands in for age, income or ethnicity, so the model discriminates without ever being given the protected attribute.",
 "Measurement bias — some groups' data is captured less accurately (for example free-text notes not in English), so information is lost before the model ever sees it.",
 "Feedback loops — the model's own decisions generate the data used to retrain it, so an unmitigated bias amplifies over successive versions.",
])
p("Bias is rarely the product of a biased developer. It is ordinary data and ordinary design choices "
  "reproducing an existing pattern — which is exactly why it must be tested for rather than assumed absent.")
h3("Implications for individuals, groups and cultures (K2, A8)")
bullets([
 "On the individual — a concrete unfair outcome: a delayed appointment, a declined application, a lower priority. It usually falls on the person least equipped to challenge it.",
 "On groups — one model decides every case identically, so an error is applied systematically to a whole group rather than randomly. Scale converts a small bias into a population-level harm.",
 "On minority groups — smaller groups are under-represented in training data, so error rates are highest exactly where the ability to contest a decision is often lowest.",
 "On language and culture — systems perform best in the language and cultural context they were trained on, a live concern in multilingual Singapore.",
 "On the organisation — regulatory exposure, contractual and legal risk, loss of trust and remediation cost, all far above the cost of testing beforehand.",
 "Across the market — algorithmic monoculture: when many organisations use the same model, the same bias and the same failure mode are entrenched everywhere at once.",
])
h3("Choosing a fairness measure (A4)")
dl([
 ("Demographic parity","Equal positive rates across groups. Suits outreach and access to opportunity; ignores genuine differences in need."),
 ("Equal opportunity","Equal true-positive rates across groups. Suits screening where a miss is the harm; needs reliable ground-truth labels."),
 ("Equalised odds","Equal true- and false-positive rates. The right target for high-stakes decisions about people, and the hardest to satisfy."),
 ("Calibration","A given score means the same thing for every group. Suits risk scoring and pricing; can coexist with unequal error rates."),
 ("Individual fairness","Similar individuals treated similarly; requires defining 'similar', which is itself a judgement."),
])
p("These measures are mathematically incompatible — you cannot satisfy them all simultaneously. "
  "Choosing the one that fits the use case, and recording why, is a governance decision rather than a "
  "technical one.")
h3("Mitigating bias, and verifying the mitigation (A5)")
bullets([
 "Examine the data for provenance, representation and proxy features; remove or generalise proxies such as postal code.",
 "Fix the fairness measure and the acceptance threshold BEFORE testing, so the result cannot be chosen after the fact.",
 "Test by group, not in aggregate — aggregate accuracy hides differential error rates.",
 "Introduce human review wherever a decision materially affects a person, and provide an explanation and an appeal route.",
 "Monitor by group after deployment: a system that was fair at launch can become unfair as the population changes.",
 "Verify by measurement — a mitigation is only real if you can show the differential error rate fell, reviewed by someone who did not build the model (NIST MEASURE 1.3). AI Verify produces that evidence.",
])
h2("The privacy-performance trade-off (A2)")
p("Richer, more granular and more identifiable data generally improves model performance, and every "
  "privacy measure removes or distorts information and so costs some accuracy. No option is free on "
  "both sides; the decision is where on the curve to sit.")
dl([
 ("Drop unneeded fields","High privacy gain, usually no performance cost. Always the first move — most training sets carry fields nobody can justify."),
 ("Generalise (age bands, districts)","Moderate to high gain, small to moderate cost. Often reduces re-identification risk AND weakens a discriminatory proxy at the same time."),
 ("Hash / pseudonymise","Limited gain — the data remains personal data and PDPA obligations continue. No modelling cost, because the identifier is not a feature."),
 ("Aggregation","High gain, moderate cost. Suits reporting and analytics rather than per-person decisions."),
 ("Differential privacy","Very high and measurable gain, but a material cost that grows with the guarantee. Rarely appropriate for a safety-critical model."),
 ("Federated learning","High gain — the data stays local — at a moderate cost plus significant engineering complexity."),
])
p("Record the decision, its reasoning and the residual risk, and have the data owner and the DPO both "
  "sign it off. An unrecorded trade-off is a decision nobody made.")
h2("Environmental impact and the energy footprint (K1, A9)")
bullets([
 "Two stages consume energy. Training is a large one-off or periodic cost; inference is small per request but recurs on every request, so for a widely used deployed system it dominates the lifetime footprint.",
 "It is not only electricity: data centres use water for cooling, and the hardware carries an embodied footprint in manufacture and disposal.",
 "Consumption depends on model size and on where and when the workload runs, because grid carbon intensity varies by region and by time of day.",
 "To estimate the footprint: energy per inference (kWh per 1,000 requests) multiplied by request volume, plus training or fine-tuning amortised as a periodic cost, plus cooling water, plus embodied carbon.",
 "The two highest-impact reductions: right-size the model (a large model used for routine work multiplies energy for no benefit), and cut unnecessary inference (cache, batch, shorten prompts and outputs, and remove AI where a rule or lookup would do).",
 "Choose the hosting region deliberately for grid carbon intensity, subject to data-residency requirements, and avoid unnecessary retraining cycles.",
 "NIST AI 600-1 lists Environmental Impacts as one of the twelve generative-AI risks, so it belongs in the risk register and not only in the sustainability report.",
])
h2("Communicating capabilities and limitations (K4)")
bullets([
 "A generative model sounds equally confident whether it is right or wrong — NIST calls confidently stated false content confabulation. Users calibrate their trust on how the system is described, so an overstated description directly causes over-reliance.",
 "Over-reliance is itself a named NIST risk (Human-AI Configuration), covering poor handoff and humans rubber-stamping model output. Design the human check; do not assume it.",
 "State what the system is NOT for. That single disclosure prevents more harm than any capability list.",
 "Publish a model card or system card: purpose, data, performance, limitations and intended use — the 'food label' transparency the MGF's trusted-development dimension calls for.",
 "Provide an obvious, easy escalation route to a human. Transparency with no route to a person is only a disclaimer.",
 "Overstated claims also carry legal and regulatory exposure, particularly in regulated sectors.",
])
h2("Championing responsible AI practice (A1, A3, A7)")
bullets([
 "Give role-specific training, not generic awareness — each group needs its own real tasks and its own concrete rules.",
 "Make the sanctioned path the easy path: an approved toolset, a fast approval route with a committed turnaround, and templates. Governance that obstructs is governance that is bypassed.",
 "Use your own organisation's near-misses; they persuade far better than external examples.",
 "Label AI-generated content and review it before publication where the audience could be harmed — content provenance is MGF dimension 7.",
 "Report results rather than intentions: commit to testing before deployment and monitoring after it, and publish what the tests found.",
 "Lead with the enabler argument — governed organisations deploy faster, because approval and evidence become routine instead of being reinvented for every project.",
])

h1("Glossary")
dl([
 ("Agent","An AI system that plans, calls tools, keeps memory and takes actions, rather than only producing text."),
 ("Agentic AI","AI systems that operate with autonomy — deciding what to do next and acting — with limited human oversight per action."),
 ("AI system inventory","The register of every AI system with its owner, purpose, data classes, model, tools, risk tier and review date. The foundational governance control."),
 ("Autonomy tier","A classification of how independently an agent operates (Tier 1 supervised to Tier 4 full autonomy) that determines the oversight obligations applied to it."),
 ("Confabulation","Confidently stated false content generated by a model; often called hallucination."),
 ("Data intermediary","Under the PDPA, an organisation that processes personal data on behalf of another; it carries the Protection and Retention Obligations."),
 ("Deployment gate","The documented evidence and approvals required before an AI system may go into production."),
 ("Guardrails","Input and output filtering at the model boundary — PII detection, jailbreak detection, content filtering. Necessary but not sufficient."),
 ("Indirect prompt injection","Malicious instructions hidden in content the agent retrieves (e-mail, documents, web pages) rather than typed by the user."),
 ("Least agency","Constraining what an agent is allowed to DECIDE, separately from what its credentials permit it to access."),
 ("Lethal trifecta","Private data access + untrusted content exposure + external communication ability, in one agent."),
 ("Model card / system card","A structured disclosure of a model or system's purpose, data, performance and limitations."),
 ("Permission ladder","A pre-execution policy that evaluates each tool call as deny, ask or allow."),
 ("Prompt injection","Input crafted so that the model treats it as instruction rather than data, redirecting the system's behaviour."),
 ("Pseudonymisation","Replacing identifiers with a consistent substitute (e.g. a hash). The data remains personal data; it is not anonymisation."),
 ("RACI","Responsible, Accountable, Consulted, Informed — a mapping of roles to controls. Exactly one Accountable per control."),
 ("Red-teaming","Adversarial testing that asks how to make a system misbehave, rather than whether it works."),
 ("Residual risk","The risk remaining after controls are applied; it must be documented and accepted by a named owner."),
 ("Risk tier","A classification (High/Medium/Low) that determines the depth of assessment, testing and oversight applied to an AI system."),
 ("Shadow AI","AI tools used within an organisation without approval, registration or oversight."),
 ("TEVV","Test, Evaluation, Verification and Validation — the assurance activities in the NIST AI RMF."),
 ("TRAQOM","The mandatory SSG course feedback survey that WSQ learners must complete for funding eligibility."),
])
# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="img":
            out+=[f"![{rest[1] or rest[0]}](courseware/assets/{rest[0]})",""]
            if rest[1]: out+=[f"*{rest[1]}*",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=_asset("tertiary-infotech-logo.png"),
                      course_logo=_asset("wsq-badge.png"), course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("1.0",C.VERSION_DATE,
  "Initial release — Learner Guide for AI Security Governance for Businesses: the full step-by-step procedure for all 16 hands-on labs across the four topics, a Framework Reference (NIST AI RMF, NIST AI 600-1, Singapore MGF and AI Verify, PDPA and the PDPC AI advisory guidelines, OWASP ASI and the CSA agentic profile), a Responsible AI Reference chapter covering ethics, bias, privacy and sustainability (K1-K5, A1-A9), embedded lab-tool screenshots and a glossary.",
  C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph()
            para.paragraph_format.left_indent=Pt(18)
            para.paragraph_format.first_line_indent=Pt(-18)
            r=para.add_run(f"{i}.  "); r.bold=True
            para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="img":
        _p=_asset(rest[0])
        if _p:
            from docx.shared import Inches as _In
            doc.add_picture(_p,width=_In(5.6))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            if rest[1]:
                cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                cr=cp.add_run(rest[1]); cr.italic=True; cr.font.size=Pt(9); cr.font.color.rgb=GREY
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
